import streamlit as st
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from io import BytesIO

# ======================
# SIDEBAR API CONFIG
# ======================

st.sidebar.title("API Configuration")

SERPAPI_KEY = st.sidebar.text_input(
    "SerpAPI Key",
    type="password"
)

OPENAI_KEY = st.sidebar.text_input(
    "OpenAI API Key",
    type="password"
)

# ======================
# UI PRINCIPALE
# ======================

st.title("SEO Article Generator")

st.write(
    "Genera articoli SEO analizzando automaticamente i competitor nella SERP."
)

keyword = st.text_input("Main keyword")

num_results = st.number_input(
    "Numero contenuti su cui fare scraping",
    min_value=1,
    max_value=10,
    value=3
)

country = st.text_input(
    "Country code (gl)",
    value="it",
    help="Inserisci il codice paese (es: it, us, gb, es, fr)"
)

language = st.text_input(
    "Language code (hl)",
    value="it",
    help="Inserisci il codice lingua (es: it, en, es, fr)"
)

generate = st.button("Genera contenuto")

# ======================
# FUNZIONI
# ======================

def get_competitors(keyword: str, num_results: int, serp_key: str, hl: str, gl: str):

    url = "https://serpapi.com/search.json"

    params = {
        "engine": "google",
        "q": keyword,
        "hl": hl,
        "gl": gl,
        "num": num_results,
        "api_key": serp_key
    }

    response = requests.get(url, params=params)

    data = response.json()

    organic = data.get("organic_results", [])

    competitors = []

    for item in organic[:num_results]:

        competitors.append({
            "title": item.get("title"),
            "link": item.get("link")
        })

    return competitors


def fetch_page(url: str):

    try:

        resp = requests.get(url, timeout=10)

        html = resp.text

        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        text = " ".join(soup.get_text().split())

        return html, text[:18000]

    except Exception:

        return "", ""


def extract_metadata(html: str):

    soup = BeautifulSoup(html, "html.parser")

    title = soup.title.string.strip() if soup.title else ""

    h1_tag = soup.find("h1")

    h1 = h1_tag.get_text(strip=True) if h1_tag else ""

    meta_desc = ""

    meta = soup.find("meta", attrs={"name": "description"})

    if meta and "content" in meta.attrs:

        meta_desc = meta["content"].strip()

    return title, h1, meta_desc


def generate_article(keyword: str, competitors: list, openai_key: str, language: str):

    client = OpenAI(api_key=openai_key)

    merged = ""

    for i, comp in enumerate(competitors, start=1):

        merged += f"""
COMPETITOR {i}

URL: {comp['link']}

TITLE: {comp['html_title']}
H1: {comp['h1']}
META: {comp['meta_desc']}

CONTENUTO:
{comp['text']}

------------------------------------
"""

    prompt = f"""
Sei un content writer SEO esperto.

Scrivi un contenuto SEO completo per la keyword:

{keyword}

IMPORTANTE:
L'articolo deve essere scritto nella stessa lingua della ricerca Google.

Language code della ricerca: {language}

Il risultato deve contenere:

TITLE TAG (max 60 caratteri)

META DESCRIPTION (max 155 caratteri)

ARTICOLO (800-1200 parole)

Regole dell'articolo:

- usa sottotitoli gerarchici
- indica SEMPRE il livello dell'heading
- usa il formato:

H2: titolo sezione

testo

H3: sottosezione

testo

- linguaggio naturale
- SEO friendly
- evita duplicazioni
- usa gli insight dei competitor senza copiarli
- scrivi tutto nella lingua indicata

COMPETITOR DATA:
{merged}

Restituisci il risultato nel seguente formato:

TITLE TAG:
...

META DESCRIPTION:
...

ARTICLE:
...
"""

    response = client.chat.completions.create(
        model="gpt-5.4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )

    content = response.choices[0].message.content

    title = ""
    meta = ""
    article = content

    if "TITLE TAG:" in content:

        parts = content.split("TITLE TAG:")[1]

        if "META DESCRIPTION:" in parts:
            title = parts.split("META DESCRIPTION:")[0].strip()

        if "ARTICLE:" in parts:
            meta = parts.split("META DESCRIPTION:")[1].split("ARTICLE:")[0].strip()
            article = parts.split("ARTICLE:")[1].strip()

    return title, meta, article


def create_word_file(title_tag, meta_description, article):

    doc = Document()

    doc.add_heading("Title Tag", level=2)
    doc.add_paragraph(title_tag)

    doc.add_heading("Meta Description", level=2)
    doc.add_paragraph(meta_description)

    doc.add_heading("Articolo", level=2)

    for line in article.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


# ======================
# ESECUZIONE
# ======================

if generate:

    if not SERPAPI_KEY or not OPENAI_KEY:

        st.error("Inserisci entrambe le API key nella sidebar")
        st.stop()

    if not keyword:

        st.error("Inserisci una keyword")
        st.stop()

    with st.spinner("Recupero competitor dalla SERP..."):

        competitors_raw = get_competitors(
            keyword,
            num_results,
            SERPAPI_KEY,
            language,
            country
        )

    if len(competitors_raw) == 0:

        st.error("Nessun competitor trovato")
        st.stop()

    competitors = []

    st.write("### Analisi contenuti competitor")

    progress = st.progress(0)
    status = st.empty()

    total = len(competitors_raw)

    for i, comp in enumerate(competitors_raw):

        status.write(f"Analizzo: {comp['link']}")

        html, text = fetch_page(comp["link"])

        html_title, h1, meta_desc = extract_metadata(html)

        competitors.append({
            **comp,
            "html_title": html_title,
            "h1": h1,
            "meta_desc": meta_desc,
            "text": text
        })

        progress.progress((i + 1) / total)

    status.empty()

    with st.spinner("Generazione contenuto con AI..."):

        title_tag, meta_description, article = generate_article(
            keyword,
            competitors,
            OPENAI_KEY,
            language
        )

    st.subheader("SEO Metadata")

    st.write("**Title Tag**")
    st.write(title_tag)

    st.write("**Meta Description**")
    st.write(meta_description)

    st.subheader("Articolo generato")

    st.write(article)

    word_file = create_word_file(
        title_tag,
        meta_description,
        article
    )

    st.download_button(
        label="Scarica documento Word",
        data=word_file,
        file_name=f"contenuto_{keyword.replace(' ','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
